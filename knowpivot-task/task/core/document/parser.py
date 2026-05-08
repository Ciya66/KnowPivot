import os
import re

from pypdf import PdfReader
from docx import Document


def extract_text_from_file(file_path: str) -> str:
    """
    从文件中提取文本内容
    :param file_path: 文件路径
    :return: 清洗后的文本内容
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在：{file_path}")

    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".pdf":
        text = extract_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_from_docx(file_path)
    elif ext in (".txt", ".md", ".json"):
        text = extract_from_txt(file_path)
    else:
        raise ValueError(f"不支持的文件类型：{ext}")

    return clean_text(text)


def extract_from_pdf(file_path: str) -> str:
    """解析 PDF，逐页提取文本"""
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        content = page.extract_text()
        if content:
            pages.append(content)
    return "\n".join(pages)


def extract_from_docx(file_path: str) -> str:
    """解析 Word 文档，提取段落和表格文本"""
    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_from_txt(file_path: str) -> str:
    """读取纯文本 / Markdown / JSON 文件"""
    with open(file_path, encoding="utf-8") as f:
        return f.read()


def clean_text(text: str) -> str:
    """清理换行、多余空格"""
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" +", " ", text)
    return text.strip()
